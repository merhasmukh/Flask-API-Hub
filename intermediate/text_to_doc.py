from flask import Flask, send_file
from docx import Document
import os

app = Flask(__name__)

def create_document_from_raw_text(raw_text, output_file="output.docx"):
    doc = Document()

    # Split the text into paragraphs (separate by newlines)
    paragraphs = raw_text.split('\n\n')

    for paragraph_text in paragraphs:
        paragraph = doc.add_paragraph()
        apply_markdown_formatting(paragraph, paragraph_text)

    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")
    return output_file

def apply_markdown_formatting(paragraph, text):
    # Simple markdown parsing (you can expand this with more rules)
    if '**' in text:  # Bold text
        text = text.replace('**', '')
        paragraph.add_run(text).bold = True
    elif '_' in text:  # Italics
        text = text.replace('_', '')
        paragraph.add_run(text).italic = True
    else:
        paragraph.add_run(text)

@app.route('/generate-doc', methods=['GET'])
def generate_doc():
    raw_text = """
    What is AI? 
    
    A Comprehensive Overview

    Artificial intelligence (AI) is a **broad field** encompassing the theory and development of computer systems able to perform tasks that normally require _human intelligence_.
    """

    output_file = "AI_Document.docx"
    
    # Generate the document
    create_document_from_raw_text(raw_text, output_file)
    
    # Send the file to the client
    if os.path.exists(output_file):
        return send_file(output_file, as_attachment=True, download_name="AI_Document.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        return "File not found", 404

if __name__ == '__main__':
    app.run(debug=True)
